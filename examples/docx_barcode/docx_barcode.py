# /// script
# requires-python = ">=3.10"
# dependencies = ["pystrich[png]", "python-docx"]
# ///
"""Insert a barcode into a Word document via get_imagedata()."""

from io import BytesIO

from docx import Document
from docx.shared import Mm

from pystrich.ean13 import EAN13Encoder

png = EAN13Encoder("5050070007664").get_imagedata()

doc = Document()
doc.add_heading("Road House (DVD)", level=1)
doc.add_paragraph("GTIN: 5050070007664")
doc.add_picture(BytesIO(png), width=Mm(64))
doc.save("docx_barcode.docx")
